import docx
from docx.shared import RGBColor
from docx.shared import Pt

# Example usage
# rtf_file = '/Users/atiab/Downloads/cache1/1101-1.rtf'
# docx_file = '/Users/atiab/Downloads/cache1/1101-1.docx'
docx_file = '/Users/atiab/Desktop/trial.docx'
# convert_rtf_to_docx2(input_file, output_file)


def extract_colored_text(docx_path, red, green, blue):
    doc = docx.Document(docx_path)
    texts = []

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.color and run.font.color.rgb == RGBColor(red, green, blue):
                texts.append(run.text)

    return texts

def print_extract(content):
    for text in content:
        print(text)


def extract_size_based(docx_file):
    doc = docx.Document(docx_file)

    for paragraph in doc.paragraphs:
        # print(f"Paragraph Style: {paragraph.style.name}")
        # Iterate through runs (segments of text)
        for run in paragraph.runs:
            print(f"Run Text: {run.text}, Run Font Size: {run.font.size}\n\n")
            print(run.font.rtl)
            # Get font size from run level if explicitly set
            font_size = run.font.size
            # print(paragraph.style.name)
            if font_size is None:
                # If font size is None at run level, check paragraph style
                font_size = paragraph.style.font.size
            if font_size is not None:
                print(f"font size is none initially. Text: {run.text}, Font Size: {font_size.pt} pt\n\n")

# print_extract(bredcrumb)

# extract_font_size_text(docx_file)